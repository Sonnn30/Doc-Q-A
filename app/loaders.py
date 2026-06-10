from pathlib import Path


def load_document(file_path: Path, file_ext: str) -> str:
    """
    Load dokumen sesuai tipe file dan kembalikan teks mentah.
    Mendukung: .pdf, .docx, .txt, .md
    """
    loaders = {
        ".pdf":  _load_pdf,
        ".docx": _load_docx,
        ".txt":  _load_text,
        ".md":   _load_text,
    }

    loader_fn = loaders.get(file_ext.lower())
    if not loader_fn:
        raise ValueError(f"Format tidak didukung: {file_ext}")

    text = loader_fn(file_path)

    if not text or not text.strip():
        raise ValueError("Dokumen kosong atau tidak dapat dibaca.")

    return text.strip()


def _load_pdf(file_path: Path) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)
    except ImportError:
        pass

    from PyPDF2 import PdfReader
    reader = PdfReader(str(file_path))
    return "\n\n".join(
        page.extract_text() for page in reader.pages if page.extract_text()
    )


def _load_docx(file_path: Path) -> str:
    import docx
    doc = docx.Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)
    return "\n\n".join(paragraphs)


def _load_text(file_path: Path) -> str:
    for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            return file_path.read_text(encoding=enc)
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError(f"Tidak dapat membaca encoding file: {file_path.name}")
